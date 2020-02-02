#Arquivo Word (docx)
# pip install docx
# pip install python-docx

#Importa o modulo 
from docx import Document
from docx.shared import Inches

#Criar um documento
document = Document()
#Adicionando um cabecalho
document.add_heading('Titulo', 0)
#Adicionar um paragrafo
p = document.add_paragraph('Isso é um paragrafo')
p.add_run('Negrito').bold = True
p.add_run('Outro texto')
p.add_run('italico').italic = True
#Cabecalho - Nivel 1
document.add_heading('Cabeçalho, Nivel 1', level=1)
#Quotas
document.add_paragraph('Quotas', style='Intense Quote')
#Lista não ordenada
document.add_paragraph(
    'primeiro item, lista não ordenada', style='List Bullet'
)
#Lista ordenada
document.add_paragraph(
    'segundo paragrafo, lista ordenada', style='List Number'
)
#Tabela
tabela = [
         {"id":1,"qtde":2,"desc":"Item 1"},
         {"id":2,"qtde":10,"desc":"Item 2"}
         ]
#criar a tabela no documento
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'id'
hdr_cells[1].text = 'Qtde'
hdr_cells[2].text = 'Produto'
for item in tabela:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item['id'])
    row_cells[1].text = str(item['qtde'])
    row_cells[2].text = str(item['desc'])

document.add_page_break()
document.save('documento01.docx')

         







